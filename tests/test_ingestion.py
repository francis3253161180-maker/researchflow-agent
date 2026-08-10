from io import BytesIO

from docx import Document
import pymupdf
from openpyxl import Workbook

from app.ingestion import _normalize, parse_upload
from app.config import Settings


def make_text_pdf(text: str) -> bytes:
    """Build a tiny standards-compliant one-page PDF for parser regression tests."""
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, value in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(value)
        content.extend(b"\nendobj\n")
    xref_at = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    content.extend(b"".join(f"{offset:010d} 00000 n \n".encode("ascii") for offset in offsets[1:]))
    content.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n".encode("ascii")
    )
    return bytes(content)


def test_markdown_parser_preserves_sections():
    parsed = parse_upload(
        "note.md",
        b"# Method\n\nThe method uses hybrid retrieval.\n\n## Evaluation\n\nThe evaluation checks grounded citations.",
    )
    assert parsed.title == "note"
    assert [block.section for block in parsed.blocks] == ["Method", "Method › Evaluation"]


def test_normalize_replaces_lone_unicode_surrogates():
    assert _normalize("valid\ud835text") == "valid?text"


def test_docx_parser_extracts_paragraphs():
    document = Document()
    document.add_heading("Paper Notes", level=1)
    document.add_paragraph("This paragraph records the experimental configuration and outcome.")
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_upload("paper-notes.docx", buffer.getvalue())
    assert parsed.filename == "paper-notes.docx"
    assert "experimental configuration" in parsed.content
    assert parsed.blocks[0].section == "Paper Notes"


def test_docx_parser_extracts_table_cells():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "ResearchFlow Agent"
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_upload("resume.docx", buffer.getvalue())

    assert "ResearchFlow Agent" in parsed.content


def test_docx_parser_preserves_multilevel_heading_path_and_table_order():
    document = Document()
    document.add_heading("3 Experiments", level=1)
    document.add_heading("3.1 Setup", level=2)
    document.add_paragraph("The setup paragraph belongs to the setup section.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Batch size"
    table.cell(0, 1).text = "8"
    document.add_heading("3.2 Results", level=2)
    document.add_paragraph("The result paragraph belongs to the result section.")
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_upload("structured.docx", buffer.getvalue())
    by_text = {block.content: block.section for block in parsed.blocks}

    assert by_text["The setup paragraph belongs to the setup section."] == "3 Experiments › 3.1 Setup"
    assert by_text["Batch size | 8"] == "3 Experiments › 3.1 Setup"
    assert by_text["The result paragraph belongs to the result section."] == "3 Experiments › 3.2 Results"


def test_xlsx_parser_keeps_sheet_and_row_range_for_citations():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "实验结果"
    sheet.append(["方法", "准确率", "备注", "设置"])
    sheet.append(["Holo", 85.43, "低参数配置", "rank-8"])
    sheet.append(["LoRA", 82.10, "baseline", "rank-8"])
    buffer = BytesIO()
    workbook.save(buffer)

    parsed = parse_upload("results.xlsx", buffer.getvalue())

    assert parsed.media_type.endswith("spreadsheetml.sheet")
    assert "Holo" in parsed.content
    assert "准确率：85.43" in parsed.content
    assert parsed.blocks[0].section == "工作表：实验结果｜行 1-3"


def test_pdf_parser_preserves_page_number_and_text():
    parsed = parse_upload("paper.pdf", make_text_pdf("ResearchFlow PDF parser regression"))

    assert parsed.filename == "paper.pdf"
    assert parsed.blocks[0].page == 1
    assert "ResearchFlow PDF parser regression" in parsed.content


def test_pdf_parser_preserves_section_path_across_pages():
    document = pymupdf.open()
    page_one = document.new_page()
    page_one.insert_text((72, 72), "1 Method", fontsize=18, fontname="hebo")
    page_one.insert_text((72, 110), "Method page one discusses calibration details.", fontsize=10)
    page_two = document.new_page()
    page_two.insert_text((72, 110), "Method page two continues the calibration discussion.", fontsize=10)
    page_three = document.new_page()
    page_three.insert_text((72, 72), "1.1 Results", fontsize=15, fontname="hebo")
    page_three.insert_text((72, 110), "Results page reports the benchmark outcome.", fontsize=10)
    payload = document.tobytes()
    document.close()

    parsed = parse_upload("structured.pdf", payload)
    by_text = {block.content: block for block in parsed.blocks}

    assert by_text["Method page one discusses calibration details."].page == 1
    assert by_text["Method page one discusses calibration details."].section == "1 Method"
    assert by_text["Method page two continues the calibration discussion."].page == 2
    assert by_text["Method page two continues the calibration discussion."].section == "1 Method"
    assert by_text["Results page reports the benchmark outcome."].page == 3
    assert by_text["Results page reports the benchmark outcome."].section == "1 Method › 1.1 Results"


def test_settings_loads_local_dotenv_without_overriding_process_env(tmp_path, monkeypatch):
    (tmp_path / ".env").write_text(
        "EMBEDDING_PROVIDER=fastembed\nRESEARCHFLOW_APP_API_KEY=from-dotenv\n",
        encoding="utf-8",
    )
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("RESEARCHFLOW_APP_API_KEY", "from-process")

    settings = Settings.from_env()

    assert settings.embedding_provider == "fastembed"
    assert settings.app_api_key == "from-process"
